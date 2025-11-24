#!/usr/bin/env python3
"""
Scanner genérico de documentos (V6 - análise de conteúdo + estrutura)

- Detecta sistema operacional (Windows / macOS / Linux)
- Verifica e instala dependências automaticamente
- Filtra binários óbvios (zip, exe, imagens etc.)
- Tenta extrair texto de qualquer arquivo textual:
    - PDF, DOCX, CSV, XLSX/XLS, TXT, MD, JSON, HTML, LOG, etc.
    - Outros formatos: tentativa de leitura como texto bruto
- Procura padrões (EMAIL, CPF, CNPJ, URL, TELEFONE, DATA, DINHEIRO, PORCENTAGEM)
- Analisa estrutura:
    - título
    - headings (seções)
    - parágrafos
    - bullets
    - estatísticas (n_palavras, n_linhas, etc.)
- Classifica tipo de documento (heurístico): cv, contrato, nota_fiscal, log, codigo_fonte, desconhecido
- Gera:
    - ./saida/_docs_index.*        → visão geral dos documentos
    - ./saida/analise_detalhada.*  → análise de conteúdo + estrutura por documento
    - ./saida/geral / por_padrao / por_arquivo → matches de padrões
    - ./saida/texto_bruto/         → texto completo de cada documento
    - ./saida/cv_extracao/         → extração estruturada de currículos (quando detectado)
"""

import os
import re
import sys
import platform
import subprocess
import argparse
from pathlib import Path
from typing import List, Dict, Any, Tuple

# =====================================
# CONFIGURAÇÕES GERAIS
# =====================================

EXCLUDE_EXTENSIONS = {
    ".zip", ".rar", ".7z", ".gz", ".tar",
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tif", ".tiff",
    ".mp3", ".wav", ".flac", ".mp4", ".mkv", ".avi", ".mov", ".wmv",
    ".exe", ".dll", ".so", ".dylib",
    ".iso",
}

MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB

# =====================================
# DETECÇÃO DE SISTEMA + DEPENDÊNCIAS
# =====================================

DEPENDENCIAS = [
    ("pandas", "pandas"),
    ("PyPDF2", "PyPDF2"),
    ("openpyxl", "openpyxl"),
    ("docx", "python-docx"),
    ("reportlab", "reportlab"),
]

def detectar_sistema() -> str:
    so = platform.system()
    if so == "Windows":
        return "Windows"
    elif so == "Darwin":
        return "macOS"
    elif so == "Linux":
        return "Linux"
    else:
        return f"Desconhecido ({so})"

def verificar_e_instalar_dependencias():
    import importlib

    so = detectar_sistema()
    print(f"\n🖥  Sistema operacional detectado: {so}")

    missing = []
    for module_name, pip_name in DEPENDENCIAS:
        try:
            importlib.import_module(module_name)
        except ImportError:
            missing.append(pip_name)

    if not missing:
        print("✅ Todas as dependências já estão instaladas.\n")
        return

    print("\n⚠  Dependências faltando:", ", ".join(missing))
    print("Tentando instalar automaticamente...")

    cmd = [sys.executable, "-m", "pip", "install"] + missing
    print(f"\n➡ Executando: {' '.join(cmd)}\n")

    try:
        subprocess.check_call(cmd)
        print("\n✅ Instalação automática concluída com sucesso!\n")
    except Exception:
        print("\n❌ Não foi possível instalar automaticamente.")
        print("Por favor, execute manualmente no terminal / cmd:")
        print(f"\n    {' '.join(cmd)}\n")
        print("Depois, rode o script novamente.")
        sys.exit(1)

verificar_e_instalar_dependencias()

import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ======================================
# PADRÕES DE INFORMAÇÃO
# ======================================

PATTERNS = {
    "EMAIL": r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+",
    "URL": r"https?://[^\s]+",
    "CPF": r"\b\d{3}\.?\d{3}\.?\d{3}-?\d{2}\b",
    "CNPJ": r"\b\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}\b",
    "TELEFONE_BR": r"\+?\d{2}\s*\(?\d{2}\)?\s*\d{4,5}-?\d{4}",
    # datas em formatos comuns: 31/12/2025, 2025-12-31 etc.
    "DATA_BR": r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",
    "DATA_ISO": r"\b\d{4}-\d{2}-\d{2}\b",
    # valores em dinheiro (R$ 1.234,56, 1.234,56, $1,234.56)
    "DINHEIRO": r"(R\$ ?\d{1,3}(\.\d{3})*,\d{2})|(\$\d{1,3}(,\d{3})*\.\d{2})",
    # porcentagens: 10%, 10,5%, 10.5 %
    "PORCENTAGEM": r"\b\d{1,3}([.,]\d+)? ?%\b",
}

COMPILED = {name: re.compile(p) for name, p in PATTERNS.items()}
FORMATS = ["csv", "xlsx", "txt", "json", "pdf", "docx"]

# ======================================
# FUNÇÕES DE UTILIDADE E LEITURA
# ======================================

def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path

def is_probably_binary(path: Path) -> bool:
    try:
        with path.open("rb") as f:
            chunk = f.read(4096)
        if not chunk:
            return False
        null_bytes = chunk.count(b"\x00")
        return null_bytes > 50
    except Exception:
        return True

def read_pdf(path: Path) -> List[Tuple[str, str]]:
    blocks = []
    try:
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                blocks.append((f"page {i}", text))
    except Exception as e:
        print(f"[ERRO] PDF {path}: {e}")
    return blocks

def read_docx(path: Path) -> List[Tuple[str, str]]:
    blocks = []
    try:
        doc = DocxDocument(str(path))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        if paras:
            text = "\n".join(paras)
            blocks.append(("docx full", text))
    except Exception as e:
        print(f"[ERRO] DOCX {path}: {e}")
    return blocks

def read_csv(path: Path) -> List[Tuple[str, str]]:
    blocks = []
    try:
        df = pd.read_csv(path, dtype=str).fillna("")
        for idx, row in df.iterrows():
            blocks.append((f"row {idx + 1}", " | ".join(map(str, row.values))))
    except Exception as e:
        print(f"[ERRO] CSV {path}: {e}")
    return blocks

def read_excel(path: Path) -> List[Tuple[str, str]]:
    blocks = []
    try:
        xls = pd.ExcelFile(path)
        for sheet in xls.sheet_names:
            df = xls.parse(sheet_name=sheet, dtype=str).fillna("")
            for idx, row in df.iterrows():
                blocks.append((f"{sheet} row {idx + 1}", " | ".join(map(str, row.values))))
    except Exception as e:
        print(f"[ERRO] Excel {path}: {e}")
    return blocks

def read_text_file(path: Path) -> List[Tuple[str, str]]:
    blocks = []
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
        if text.strip():
            blocks.append(("text full", text))
    except Exception as e:
        print(f"[ERRO] TXT genérico {path}: {e}")
    return blocks

def extract_blocks_generic(path: Path) -> List[Tuple[str, str]]:
    ext = path.suffix.lower()

    if ext in EXCLUDE_EXTENSIONS:
        return []

    try:
        if path.stat().st_size > MAX_FILE_SIZE:
            print(f"[AVISO] {path} ignorado (arquivo muito grande).")
            return []
    except OSError:
        return []

    # arquivos com leitores dedicados não passam pelo detector binário
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    if ext == ".csv":
        return read_csv(path)
    if ext in [".xlsx", ".xls"]:
        return read_excel(path)

    text_like_exts = {".txt", ".md", ".json", ".html", ".htm", ".log", ".ini", ".cfg"}
    if ext in text_like_exts:
        return read_text_file(path)

    if is_probably_binary(path):
        print(f"[AVISO] {path} parece binário, ignorando.")
        return []

    return read_text_file(path)

def find_patterns(text: str):
    results = []
    for pname, regex in COMPILED.items():
        for m in regex.finditer(text):
            value = m.group(0)
            context = text[max(0, m.start() - 40): m.end() + 40].replace("\n", " ")
            results.append((pname, value, context))
    return results

# ======================================
# ANÁLISE DE ESTRUTURA E CONTEÚDO
# ======================================

STOPWORDS_PT = {
    "de","da","do","para","por","em","um","uma","e","a","o","os","as",
    "que","com","no","na","nas","nos","se","ao","à","às","aos"
}
STOPWORDS_EN = {
    "the","and","of","in","to","for","on","a","an","is","it","that","with","as","at","by"
}

def sample_items(items: List[str], limit: int = 8) -> List[str]:
    if not items:
        return []
    if len(items) <= limit:
        return items
    if limit <= 1:
        return items[:1]
    span = len(items) - 1
    positions = [round(span * i / (limit - 1)) for i in range(limit)]
    sampled = []
    seen = set()
    for idx in positions:
        idx = max(0, min(idx, len(items) - 1))
        value = items[idx]
        lower = value.lower()
        if lower in seen:
            continue
        sampled.append(value)
        seen.add(lower)
    return sampled

def guess_language(text: str) -> str:
    lower = text.lower()
    count_pt = sum(lower.count(w + " ") for w in STOPWORDS_PT)
    count_en = sum(lower.count(" " + w + " ") for w in STOPWORDS_EN)
    if count_pt > count_en and count_pt > 3:
        return "pt"
    if count_en > count_pt and count_en > 3:
        return "en"
    return "desconhecido"

def split_paragraphs(text: str) -> List[str]:
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    return paras

def detect_headings(lines: List[str]) -> List[str]:
    headings = []
    seen = set()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        normalized = re.sub(r"\s+", " ", stripped)
        if len(normalized) > 140:
            continue
        if normalized.count("|") >= 2:
            continue
        if sum(1 for c in normalized if c.isalpha()) < 3:
            continue

        candidate = normalized.rstrip(":").strip()
        lower = candidate.lower()
        if lower in seen:
            continue

        words = candidate.split()
        if not words:
            continue
        first_token = words[0]
        number_pattern = r"^(\d+(\.\d+)*|\(?[IVXLC]+\)?|[A-Z]\.)$"
        is_numbered = bool(re.match(number_pattern, first_token))
        uppercase_ratio = sum(1 for w in words if w.isupper()) / len(words)
        ends_with_colon = stripped.endswith(":")
        camel_case = all(len(w) > 0 and w[0].isupper() for w in words if any(ch.isalpha() for ch in w))

        looks_heading = False
        if is_numbered and len(words) <= 14:
            looks_heading = True
        elif uppercase_ratio >= 0.75 and len(words) <= 10:
            looks_heading = True
        elif ends_with_colon and len(words) <= 9:
            looks_heading = True
        elif camel_case and len(words) <= 6:
            looks_heading = True

        if looks_heading:
            headings.append(candidate)
            seen.add(lower)

    return headings

def detect_bullets(lines: List[str]) -> List[str]:
    bullets = []
    for line in lines:
        if re.match(r"^\s*[-*•]\s+", line):
            bullets.append(line.strip())
    return bullets

def top_words(text: str, n=10) -> str:
    words = re.findall(r"\b\w+\b", text.lower())
    freq = {}
    for w in words:
        if w in STOPWORDS_PT or w in STOPWORDS_EN:
            continue
        if len(w) <= 2:
            continue
        freq[w] = freq.get(w, 0) + 1
    top = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:n]
    return ", ".join(f"{w}({c})" for w, c in top)

def classify_doc_type(text: str) -> str:
    lower = text.lower()

    cv_keywords = [
        "currículo", "curriculo", "curriculum",
        "experiência profissional", "experiencia profissional",
        "formação acadêmica", "formacao academica",
        "education", "work experience", "skills", "habilidades",
        "resumo profissional", "objetivo profissional"
    ]
    if sum(1 for k in cv_keywords if k in lower) >= 2:
        return "cv"

    if any(k in lower for k in ["contrato", "cláusula", "clausula", "contratante", "contratada"]):
        return "contrato"

    if any(k in lower for k in ["nota fiscal", "nfe", "nf-e", "cnpj emitente", "danfe"]):
        return "nota_fiscal"

    if any(k in lower for k in ["exception", "stack trace", "log level", "timestamp", "ERROR", "WARN"]):
        return "log"

    if any(sym in text for sym in [";}", "public class", "def ", "function ", "#include", "console.log", "using System"]):
        return "codigo_fonte"

    return "desconhecido"

def analyze_text_structure(file_name: str, text: str) -> Dict[str, Any]:
    lines = [l.rstrip() for l in text.splitlines()]
    paras = split_paragraphs(text)

    headings = detect_headings(lines)
    bullets = detect_bullets(lines)

    n_chars = len(text)
    n_words = len(re.findall(r"\b\w+\b", text))
    n_lines = len(lines)
    n_paragraphs = len(paras)
    avg_words_per_paragraph = n_words / n_paragraphs if n_paragraphs else 0

    title = ""
    for l in lines:
        if l.strip():
            title = l.strip()
            break

    lang = guess_language(text)
    doc_type = classify_doc_type(text)

    topw = top_words(text, n=12)

    n_tables_guess = 0
    for l in lines:
        if ";" in l or "\t" in l:
            n_tables_guess += 1

    resumo_rapido = " ".join(paras[0].split()[:40]) if paras else ""

    return {
        "file_name": file_name,
        "doc_type": doc_type,
        "language_guess": lang,
        "title": title[:200],
        "n_chars": n_chars,
        "n_words": n_words,
        "n_lines": n_lines,
        "n_paragraphs": n_paragraphs,
        "avg_words_per_paragraph": round(avg_words_per_paragraph, 2),
        "n_headings": len(headings),
        "headings_sample": " | ".join(sample_items(headings, 8)),
        "n_bullets": len(bullets),
        "bullets_sample": " | ".join(sample_items(bullets, 8)),
        "n_tables_guess": n_tables_guess,
        "top_words": topw,
        "resumo_rapido": resumo_rapido,
    }

# ======================================
# VARREDURA PRINCIPAL
# ======================================

def scan_directory(base: Path):
    """
    Retorna:
      - df_matches: DataFrame com todos os matches de padrões
      - raw_texts: dict[file_name] = texto completo
      - docs_index: lista com visão geral por documento
      - docs_analysis: lista com análise de conteúdo + estrutura por documento
    """
    all_results: List[Dict[str, Any]] = []
    raw_texts: Dict[str, str] = {}
    docs_index: List[Dict[str, Any]] = []
    docs_analysis: List[Dict[str, Any]] = []

    for root, _, files in os.walk(base):
        for fname in files:
            path = Path(root) / fname

            blocks = extract_blocks_generic(path)
            if not blocks:
                continue

            print(f"[LENDO] {path}")

            full_text = "\n\n".join(text for _, text in blocks)
            if not full_text.strip():
                continue

            raw_texts[path.name] = full_text

            ext = path.suffix.lower()
            size = path.stat().st_size if path.exists() else 0
            doc_type = classify_doc_type(full_text)

            docs_index.append({
                "file_name": path.name,
                "file_path": str(path),
                "file_type_ext": ext.replace(".", ""),
                "file_size_bytes": size,
                "doc_type": doc_type,
            })

            # análise detalhada de estrutura/conteúdo
            analysis = analyze_text_structure(path.name, full_text)
            docs_analysis.append(analysis)

            for location, text in blocks:
                matches = find_patterns(text)
                for pname, value, context in matches:
                    all_results.append({
                        "file_name": path.name,
                        "file_path": str(path),
                        "file_type": ext.replace(".", ""),
                        "location": location,
                        "pattern": pname,
                        "value": value,
                        "context": context,
                    })

    df_matches = pd.DataFrame(all_results)
    return df_matches, raw_texts, docs_index, docs_analysis

# ======================================
# EXPORTAÇÃO
# ======================================

def export_formats(df: pd.DataFrame, out_dir: Path, base_name: str):
    ensure_dir(out_dir)

    has_pattern_schema = {"pattern", "value", "file_name", "location"}.issubset(df.columns)

    for fmt in FORMATS:
        out_file = out_dir / f"{base_name}.{fmt}"

        if fmt == "csv":
            df.to_csv(out_file, index=False, encoding="utf-8-sig")

        elif fmt == "xlsx":
            df.to_excel(out_file, index=False)

        elif fmt == "json":
            df.to_json(out_file, orient="records", indent=4, force_ascii=False)

        elif fmt == "txt":
            with open(out_file, "w", encoding="utf-8") as f:
                for _, row in df.iterrows():
                    if has_pattern_schema:
                        line = f"[{row['pattern']}] {row['value']}  ({row['file_name']} - {row['location']})"
                    else:
                        pairs = [f"{col}={row[col]}" for col in df.columns]
                        line = " | ".join(pairs)
                    f.write(line + "\n")

        elif fmt == "docx":
            doc = DocxDocument()
            if has_pattern_schema:
                doc.add_heading("Resultados da varredura", level=1)
                for _, row in df.iterrows():
                    p = doc.add_paragraph()
                    p.add_run(f"[{row['pattern']}] ").bold = True
                    p.add_run(f"{row['value']}  ")
                    p.add_run(f"({row['file_name']} - {row['location']})").italic = True
            else:
                doc.add_heading("Resultados (genérico)", level=1)
                for _, row in df.iterrows():
                    p = doc.add_paragraph()
                    for col in df.columns:
                        p.add_run(f"{col}: ").bold = True
                        p.add_run(f"{row[col]}  ")
                    doc.add_paragraph("")
            doc.save(out_file)

        elif fmt == "pdf":
            c = canvas.Canvas(str(out_file), pagesize=letter)
            y = 750
            for _, row in df.iterrows():
                if has_pattern_schema:
                    text = f"[{row['pattern']}] {row['value']} ({row['file_name']} - {row['location']})"
                else:
                    parts = [f"{col}={row[col]}" for col in df.columns]
                    text = " | ".join(parts)
                c.drawString(40, y, text[:130])
                y -= 14
                if y < 60:
                    c.showPage()
                    y = 750
            c.save()

        print(f"[OK] Exportado: {out_file}")

def export_raw_texts(raw_texts: Dict[str, str], saida_base: Path):
    base = ensure_dir(saida_base / "texto_bruto")
    for fname, text in raw_texts.items():
        safe_name = fname.replace("/", "_")
        out_file = base / f"{safe_name}.txt"
        with open(out_file, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"[OK] Texto bruto salvo: {out_file}")

def export_docs_index(docs_index: List[Dict[str, Any]], saida_base: Path):
    if not docs_index:
        return
    df_idx = pd.DataFrame(docs_index)
    export_formats(df_idx, saida_base, "_docs_index")

def export_docs_analysis(docs_analysis: List[Dict[str, Any]], saida_base: Path):
    if not docs_analysis:
        return
    df_an = pd.DataFrame(docs_analysis)
    export_formats(df_an, saida_base, "analise_detalhada")

def export_all(df_matches: pd.DataFrame,
               raw_texts: Dict[str, str],
               docs_index: List[Dict[str, Any]],
               docs_analysis: List[Dict[str, Any]],
               saida_base: Path):
    export_docs_index(docs_index, saida_base)
    export_docs_analysis(docs_analysis, saida_base)

    if not df_matches.empty:
        export_formats(df_matches, saida_base / "geral", "todos_resultados")

        for pattern in df_matches["pattern"].unique():
            df_pat = df_matches[df_matches["pattern"] == pattern]
            export_formats(df_pat, saida_base / "por_padrao" / pattern, f"resultados_{pattern}")

        for fname in df_matches["file_name"].unique():
            df_file = df_matches[df_matches["file_name"] == fname]
            safe_name = fname.replace(".", "_")
            export_formats(df_file, saida_base / "por_arquivo" / safe_name, "resultados")

    export_raw_texts(raw_texts, saida_base)

# ======================================
# EXTRATOR DE CV (mesmo esquema de antes)
# ======================================

REGEX_CIDADE_UF = re.compile(
    r"\b([A-ZÁÉÍÓÚÂÊÔÃÕÇ][a-záéíóúâêôãõç]+(?:\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇa-záéíóúâêôãõç]+)*)\s*-\s*([A-Z]{2})\b"
)

def is_probable_cv(text: str) -> bool:
    lower = text.lower()
    keywords = [
        "currículo", "curriculo", "curriculum",
        "experiência profissional", "experiencia profissional",
        "formação acadêmica", "formacao academica",
        "education", "work experience", "skills", "habilidades",
        "resumo profissional", "objetivo profissional"
    ]
    hits = sum(1 for k in keywords if k in lower)
    return hits >= 2

def guess_name(text: str) -> str:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    ignore_prefixes = ["currículo", "curriculo", "curriculum", "resume"]
    for line in lines[:10]:
        lower = line.lower()
        if any(lower.startswith(p) for p in ignore_prefixes):
            continue
        if "@" in line or "http" in lower:
            continue
        if len(line.split()) >= 2 and any(c.isalpha() for c in line):
            return line
    return ""

def extract_cv_structured(file_name: str, text: str) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "file_name": file_name,
        "cv_detectado": False,
        "nome": "",
        "emails": "",
        "telefones": "",
        "linkedin": "",
        "github": "",
        "outras_urls": "",
        "localizacao": "",
        "resumo_inicial": "",
    }

    if not is_probable_cv(text):
        return result

    result["cv_detectado"] = True
    result["nome"] = guess_name(text)

    emails = set(COMPILED["EMAIL"].findall(text))
    urls = set(COMPILED["URL"].findall(text))
    telefones = set(COMPILED["TELEFONE_BR"].findall(text))

    linkedin = {u for u in urls if "linkedin.com" in u.lower()}
    github = {u for u in urls if "github.com" in u.lower()}
    outras_urls = urls - linkedin - github

    result["emails"] = "; ".join(sorted(emails))
    result["telefones"] = "; ".join(sorted(telefones))
    result["linkedin"] = "; ".join(sorted(linkedin))
    result["github"] = "; ".join(sorted(github))
    result["outras_urls"] = "; ".join(sorted(outras_urls))

    m = REGEX_CIDADE_UF.search(text)
    if m:
        result["localizacao"] = f"{m.group(1)} - {m.group(2)}"

    lines = [l.strip() for l in text.splitlines() if l.strip()]
    resumo = []
    for line in lines:
        if "@" in line or "http" in line.lower():
            continue
        resumo.append(line)
        if len(resumo) >= 5:
            break
    result["resumo_inicial"] = " | ".join(resumo)

    return result

def export_cv_summary(raw_texts: Dict[str, str], saida_base: Path):
    rows = []
    for fname, text in raw_texts.items():
        info = extract_cv_structured(fname, text)
        if info["cv_detectado"]:
            rows.append(info)

    if not rows:
        print("\nℹ Nenhum currículo identificado pelos padrões heurísticos.")
        return

    df_cv = pd.DataFrame(rows)
    out_dir = ensure_dir(saida_base / "cv_extracao")
    export_formats(df_cv, out_dir, "curriculos")
    print(f"\n[OK] Extração de CVs salva em: {out_dir}")

# ======================================
# INTERFACE / CLI
# ======================================

def parse_args():
    parser = argparse.ArgumentParser(description="Scanner genérico de documentos (V6)")
    parser.add_argument(
        "--path",
        help="Pasta com documentos a processar (ignora prompt se informado).",
    )
    parser.add_argument(
        "--auto",
        action="store_true",
        help="Usa './documentos' automaticamente (sem prompt).",
    )
    return parser.parse_args()


def main():
    args = parse_args()
    interactive = not (args.auto or args.path)

    if interactive:
        print("=" * 60)
        print("  VARREDOR GENÉRICO DE DOCUMENTOS + ANÁLISE ESTRUTURAL - V6")
        print("=" * 60)
        print("\nEste programa tenta extrair TEXTO de qualquer arquivo que pareça documento")
        print("(PDF, DOCX, CSV, Excel, TXT, JSON, HTML, LOG, etc.) e:")
        print("  - procura padrões (EMAIL, CPF, CNPJ, URL, TELEFONE, DATAS, DINHEIRO, %)")
        print("  - analisa estrutura (título, headings, parágrafos, bullets, estatísticas)")
        print("  - classifica tipo (cv, contrato, nota_fiscal, log, codigo_fonte, ...)")
        print("  - salva texto completo e análises em ./saida/\n")
        print("Passo 1: Escolha uma pasta com os documentos.")
        print("Se não souber, coloque os arquivos em uma pasta chamada 'documentos'")
        print("no mesmo lugar deste programa e apenas aperte ENTER.\n")

    if args.path:
        base_dir = Path(args.path).expanduser()
    else:
        if interactive:
            try:
                base_input = input("Digite o caminho da pasta com documentos (ou ENTER para usar './documentos'): ").strip()
            except EOFError:
                print("\n[INFO] Nenhuma entrada detectada (modo não interativo). Usando './documentos'.")
                base_input = ""
        else:
            base_input = ""

        base_dir = Path(base_input) if base_input else Path("./documentos")

    if not base_dir.exists():
        print(f"\n[AVISO] A pasta '{base_dir}' não existe.")
        print("Criando essa pasta agora...")
        base_dir.mkdir(parents=True, exist_ok=True)
        print(f"Pasta criada: {base_dir.resolve()}")
        print("\n👉 Agora coloque seus arquivos dentro dessa pasta")
        print("e rode o programa novamente.")
        return

    print(f"\n📂 Pasta selecionada: {base_dir.resolve()}")
    print("Procurando arquivos de qualquer tipo que contenham texto...")

    saida_dir = Path("./saida")
    ensure_dir(saida_dir)

    df_matches, raw_texts, docs_index, docs_analysis = scan_directory(base_dir)

    export_all(df_matches, raw_texts, docs_index, docs_analysis, saida_dir)
    export_cv_summary(raw_texts, saida_dir)

    if df_matches.empty:
        print("\n⚠ Nenhum padrão (EMAIL/CPF/CNPJ/URL/TELEFONE/DATA/DINHEIRO/%) encontrado.")
    else:
        print(f"\n✅ Itens encontrados: {len(df_matches)} (matches de padrões).")

    print("\n🎉 Pronto!")
    print(f"Abra a pasta: {saida_dir.resolve()}")
    print("Você verá, por exemplo:")
    print("  - '_docs_index.*'    → visão geral por documento (tipo, tamanho, extensão)")
    print("  - 'analise_detalhada.*' → estrutura + estatísticas de conteúdo")
    print("  - 'geral' / 'por_padrao' / 'por_arquivo' → matches detalhados")
    print("  - 'texto_bruto'      → TEXTO COMPLETO por documento (.txt)")
    print("  - 'cv_extracao'      → curriculos.* com extração estruturada (se houver)\n")


if __name__ == "__main__":
    main()
